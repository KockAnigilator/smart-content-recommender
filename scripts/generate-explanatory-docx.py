from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "docs" / "explanatory-note-gost-7.32.md"
TARGET_DOCX = ROOT / "docs" / "explanatory-note-gost-7.32.docx"


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level <= 2 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(14)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    p.add_run(f"- {text}")


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.add_run(text)


def build_docx_from_markdown(md_text: str) -> Document:
    doc = Document()
    configure_document(doc)

    for raw_line in md_text.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("---"):
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
            continue
        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            continue
        if line.startswith("```"):
            continue
        add_paragraph(doc, line)

    return doc


def main() -> None:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Source markdown not found: {SOURCE_MD}")

    text = SOURCE_MD.read_text(encoding="utf-8")
    doc = build_docx_from_markdown(text)
    doc.save(TARGET_DOCX)
    print(f"Generated: {TARGET_DOCX}")


if __name__ == "__main__":
    main()
